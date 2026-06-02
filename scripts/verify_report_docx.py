#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Basic DOCX validation for generated lab reports."""
from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document


PLACEHOLDER_PATTERNS = [
    re.compile(r"【[^】]+】"),
    re.compile(r"\[[A-Z_ ]{3,}\]"),
    re.compile(r"TODO", re.I),
]


def count_placeholders(text: str) -> int:
    return sum(len(pattern.findall(text)) for pattern in PLACEHOLDER_PATTERNS)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", help="Generated DOCX report")
    parser.add_argument("--json", action="store_true", help="Print JSON")
    args = parser.parse_args()

    path = Path(args.docx)
    if not path.exists():
        raise SystemExit(f"DOCX not found: {path}")

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    all_text = "\n".join(paragraphs)

    with zipfile.ZipFile(path) as zf:
        images = [n for n in zf.namelist() if n.startswith("word/media/")]

    headings = []
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading") and p.text.strip():
            headings.append(p.text.strip())

    result = {
        "path": str(path),
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "images": len(images),
        "placeholders": count_placeholders(all_text),
        "headings": headings[:50],
    }

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        for key, value in result.items():
            if key == "headings":
                print("headings:")
                for heading in value:
                    print(f"  - {heading}")
            else:
                print(f"{key}: {value}")


if __name__ == "__main__":
    main()
